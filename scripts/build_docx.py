#!/usr/bin/env python
"""Render PROJECT.md as a Word document.

    python3 scripts/build_docx.py                  # -> Palimpsest-Project.docx
    python3 scripts/build_docx.py --in docs/05-esl.md --out ESL-Study.docx

PROJECT.md is the source of truth and this is a VIEW of it, regenerated rather than
maintained. Keeping a second hand-edited copy of a document full of measured numbers is how
the numbers drift apart, and this project has already been bitten once by a figure that was
true when written and false when read.

Needs python-docx (present on the system interpreter here, not in the project venv, which is
deliberately kept to what the detector itself requires):

    python3 -m pip install python-docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
ACCENT = RGBColor(0x8B, 0x2E, 0x1F)      # the interface's flag colour
LINKCOL = RGBColor(0x1F, 0x56, 0x73)

#: **bold** / *italic* / `code` / [text](target), matched in one pass so nested runs do not
#: get double-processed.
INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*\n]+?\*(?!\*)|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")


def add_page_numbers(section) -> None:
    """Footer page number, as a real Word field so it repaginates with the document.

    A literal number would be wrong the moment anything reflows; PAGE is computed by Word.
    """
    par = section.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    for kind, text in (("begin", None), ("instrText", " PAGE "), ("end", None)):
        el = OxmlElement(f"w:fld{'Char' if kind != 'instrText' else ''}"
                         if kind != "instrText" else "w:instrText")
        if kind == "instrText":
            el.set(qn("xml:space"), "preserve")
            el.text = text
        else:
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)


def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_inline(par, text: str, *, base_size: float = 10.5, color: RGBColor = INK) -> None:
    """Write markdown inline formatting into a paragraph as real Word runs."""
    for part in INLINE.split(text):
        if not part:
            continue
        run = par.add_run()
        run.font.size = Pt(base_size)
        run.font.color.rgb = color
        if part.startswith("**") and part.endswith("**"):
            run.text, run.bold = part[2:-2], True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Menlo"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = ACCENT
        elif (m := re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", part)):
            # A link label is often itself a code span -- [`scripts/foo.py`](scripts/foo.py).
            # The label must be unwrapped, or the backticks reach the page as literal
            # characters; that is the only markdown that survived the first build.
            label = m.group(1)
            if label.startswith("`") and label.endswith("`") and len(label) > 2:
                label = label[1:-1]
                run.font.name = "Menlo"
                run.font.size = Pt(base_size - 1)
            run.text = label
            run.font.color.rgb = LINKCOL
            run.underline = True
        elif part.startswith("*") and part.endswith("*"):
            run.text, run.italic = part[1:-1], True
        else:
            run.text = part


def is_table_row(line: str) -> bool:
    return line.startswith("|") and line.rstrip().endswith("|")


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--in", dest="src", default="PROJECT.md")
    ap.add_argument("--out", dest="out", default="Palimpsest-Project.docx")
    args = ap.parse_args()

    src = (ROOT / args.src) if not Path(args.src).is_absolute() else Path(args.src)
    lines = src.read_text(encoding="utf-8").splitlines()

    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)
        add_page_numbers(s)
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.18

    i, first_heading = 0, True
    while i < len(lines):
        line = lines[i].rstrip()

        # -- fenced code -------------------------------------------------------
        if line.startswith("```"):
            i += 1
            buf: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run("\n".join(buf))
            r.font.name = "Menlo"
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # -- tables ------------------------------------------------------------
        if is_table_row(line) and i + 1 < len(lines) and set(lines[i + 1]) <= set("|-: "):
            header = split_row(line)
            i += 2
            body: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                body.append(split_row(lines[i]))
                i += 1
            width = max([len(header)] + [len(r) for r in body])
            t = doc.add_table(rows=1, cols=width)
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, txt in enumerate(header + [""] * (width - len(header))):
                cell = t.rows[0].cells[c]
                cell.text = ""
                shade(cell, "F0EDE6")
                par = cell.paragraphs[0]
                par.paragraph_format.space_after = Pt(2)
                add_inline(par, f"**{txt}**" if txt and "**" not in txt else txt, base_size=9)
            for row in body:
                cells = t.add_row().cells
                for c, txt in enumerate(row + [""] * (width - len(row))):
                    cells[c].text = ""
                    par = cells[c].paragraphs[0]
                    par.paragraph_format.space_after = Pt(2)
                    add_inline(par, txt, base_size=9)
            doc.add_paragraph()
            continue

        # -- headings ----------------------------------------------------------
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            if first_heading and level == 1:
                # Title block, matching the document's own opening.
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.size, r.bold, r.font.name = Pt(22), True, "Georgia"
                r.font.color.rgb = INK
                first_heading = False
                i += 1
                continue
            h = doc.add_heading(level=min(level, 4))
            h.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
            h.paragraph_format.space_after = Pt(4)
            hr = h.add_run(text)
            hr.font.name = "Georgia"
            hr.font.color.rgb = INK if level <= 2 else MUTED
            hr.font.size = Pt({1: 18, 2: 14.5, 3: 12, 4: 11}.get(level, 11))
            hr.bold = True
            i += 1
            continue

        # -- rules, blank ------------------------------------------------------
        if line.strip() in {"---", "***", "___"}:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            pr = p._p.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "D8D2C6")
            bdr.append(bottom)
            pr.append(bdr)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue

        # -- lists -------------------------------------------------------------
        if re.match(r"^\s*[-*+]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, re.sub(r"^\s*[-*+]\s+", "", line))
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, re.sub(r"^\s*\d+\.\s+", "", line))
            i += 1
            continue

        # -- blockquote --------------------------------------------------------
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, line.lstrip("> ").strip(), color=MUTED)
            i += 1
            continue

        # -- paragraph: join wrapped lines --------------------------------------
        buf = [line]
        i += 1
        while (i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", "```", ">"))
               and not re.match(r"^\s*([-*+]|\d+\.)\s+", lines[i])
               and lines[i].strip() not in {"---", "***", "___"}):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(buf))

    out = (ROOT / args.out) if not Path(args.out).is_absolute() else Path(args.out)
    doc.save(out)
    print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
