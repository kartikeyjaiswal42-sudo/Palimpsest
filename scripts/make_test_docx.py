#!/usr/bin/env python
"""Write a .docx with a real word-processor library, for the upload check to read.

    python3 scripts/make_test_docx.py

The fixtures in ``scripts/verify_upload.cjs`` are ZIPs that script builds itself, so on their
own they only prove the reader handles what *that script* emits. This one is produced by
python-docx, which writes the same package structure Word does -- section properties, styles,
relationships, a `word/document.xml` with the run/paragraph nesting a word processor actually
produces rather than the minimum the format allows. It is the difference between "the parser
agrees with my idea of a .docx" and "the parser opens a real one".

Written to the temp directory rather than committed, for the reason the check's own header
gives: a binary in the repository cannot be reviewed in a diff. The check skips this case
when the file is absent, so the suite still runs on a machine without python-docx.

Needs python-docx on the system interpreter (the project venv is deliberately kept to what
the detector itself requires):

    python3 -m pip install python-docx
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

OUT = Path(tempfile.gettempdir()) / "palimpsest-real.docx"

# Matched to what verify_upload.cjs looks for: "bus transfers" and "shoebox".
PARAGRAPHS = [
    "My grandmother kept her bus transfers. All of them, going back years, in a shoebox "
    "under the sink with the shoe polish.",
    "Nobody could explain it. My mother said she was just like that. My uncle said it was "
    "the war, which is what he says about everything.",
    "So I started keeping things too. Ticket stubs, mostly, and the receipt from the diner "
    "where my father told me he was moving out.",
]


def main() -> int:
    try:
        from docx import Document
    except ImportError:
        print(
            "python-docx is not installed on this interpreter.\n"
            "  python3 -m pip install python-docx\n"
            "The upload check skips the real-word-processor case without it.",
            file=sys.stderr,
        )
        return 1

    doc = Document()
    doc.add_heading("Personal statement", level=1)
    for text in PARAGRAPHS:
        doc.add_paragraph(text)
    # A paragraph carrying several runs, which is what a word processor produces the moment
    # anything is italicised -- the reader has to join runs back into one sentence, and a
    # naive one that treats each run as a paragraph fails exactly here.
    par = doc.add_paragraph("I do not know yet what I want to ")
    par.add_run("study").italic = True
    par.add_run(". I know I want to work with things people left behind.")
    doc.save(OUT)

    print(f"wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
